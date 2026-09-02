"""Structured Word close-report from live controller figures. No LLM numbers."""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from serialize import paise_to_rupees


NAVY = RGBColor(0x0B, 0x3A, 0x6E)
BLUE = RGBColor(0x0D, 0x4F, 0xFF)
MUTED = RGBColor(0x64, 0x74, 0x8B)


def _set_run(run, size=11, bold=False, color=NAVY, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def _shade_cell(cell, fill: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_text(cell, text, *, bold=False, color=NAVY, size=10, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text if text is not None else "—"))
    _set_run(run, size=size, bold=bold, color=color)
    if fill:
        _shade_cell(cell, fill)


def _add_heading(doc, text, size=18):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    _set_run(run, size=size, bold=True, color=NAVY)


def _add_body(doc, text, size=11):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    _set_run(run, size=size, bold=False, color=RGBColor(0x33, 0x41, 0x55))


def _kpi_table(doc, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    _set_cell_text(table.rows[0].cells[0], "Metric", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), fill="0B3A6E")
    _set_cell_text(table.rows[0].cells[1], "Value", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), fill="0B3A6E")
    for index, (label, value) in enumerate(rows, start=1):
        fill = "F8FBFF" if index % 2 else "FFFFFF"
        _set_cell_text(table.rows[index].cells[0], label, fill=fill)
        _set_cell_text(table.rows[index].cells[1], value, bold=True, fill=fill)


def _inr_paise(value):
    return f"₹{paise_to_rupees(value):,.2f}"


def _inr_rupees(value):
    try:
        return f"₹{float(value or 0):,.2f}"
    except (TypeError, ValueError):
        return "—"


def build_word_report(*, metrics, analytics, exceptions, cash, tax, audit, batch_meta=None) -> bytes:
    metrics = metrics or {}
    analytics = analytics or {}
    exceptions = exceptions or []
    cash = cash or {}
    tax = tax or {}
    audit = audit or []
    batch = batch_meta or metrics.get("batch") or {}
    generated = datetime.now(timezone.utc).astimezone().strftime("%d %b %Y, %H:%M")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = kicker.add_run("RAZOR-AI  ·  AI FINANCE CONTROLLER")
    _set_run(run, size=10, bold=True, color=BLUE)

    title = doc.add_paragraph()
    run = title.add_run("Books close report")
    _set_run(run, size=26, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    run = sub.add_run(
        f"Generated {generated}  ·  Batch {batch.get('batch_id') or '—'}  ·  Source {batch.get('source') or 'live session'}"
    )
    _set_run(run, size=10, color=MUTED)

    _add_body(
        doc,
        "This document is assembled from the live reconciliation run. Match rate, cash, GST and "
        "exceptions come from deterministic arithmetic. Gemini is not used to invent a UTR, settlement, or cash figure.",
    )

    match = metrics.get("match_rate")
    match_label = f"{(match * 100):.1f}%" if match is not None else "—"
    _add_heading(doc, "1. Executive summary", 16)
    _kpi_table(doc, [
        ("Payments in batch", f"{int(metrics.get('total_records') or analytics.get('total_orders') or 0):,}"),
        ("Match rate", match_label),
        ("GMV", _inr_rupees(analytics.get("total_earnings"))),
        ("Amount matched", _inr_paise(metrics.get("amount_reconciled"))),
        ("Open exceptions", str(
            metrics.get("unresolved_exceptions")
            if metrics.get("unresolved_exceptions") is not None
            else (metrics.get("exceptions") or len(exceptions))
        )),
        ("Amount at risk", _inr_paise(metrics.get("amount_at_risk"))),
        ("Net settlement", _inr_rupees(analytics.get("net_settlement"))),
        ("GST collected", _inr_rupees(analytics.get("total_tax") or tax.get("actual_gst_rupees"))),
    ])

    _add_heading(doc, "2. Cash position", 16)
    _kpi_table(doc, [
        ("Available now", f"₹{float(cash.get('available_rupees') or 0):,.2f}"),
        ("In transit (T+2)", f"₹{float(cash.get('in_transit_rupees') or 0):,.2f}"),
        ("Blocked by exceptions", f"₹{float(cash.get('blocked_rupees') or 0):,.2f}"),
        ("Expected next 7 days", f"₹{float(cash.get('expected_7d_rupees') or 0):,.2f}"),
    ])
    _add_body(doc, "Captured is not settled. Available cash already excludes recorded synthetic withdrawals.")

    _add_heading(doc, "3. GST ledger", 16)
    expected = float(tax.get("expected_gst_rupees") or 0)
    actual = float(tax.get("actual_gst_rupees") or tax.get("gst_collected_rupees") or 0)
    _kpi_table(doc, [
        ("Rate", tax.get("rate") or "18% of processing fee"),
        ("Expected GST", f"₹{expected:,.2f}"),
        ("GST collected", f"₹{actual:,.2f}"),
        ("Difference", f"₹{(actual - expected):,.2f}"),
        ("Mismatched lines", str(tax.get("mismatched_lines") or 0)),
    ])

    _add_heading(doc, "4. Exception register", 16)
    open_rows = [item for item in exceptions if item.get("open", True)][:40]
    if not open_rows:
        _add_body(doc, "No open exceptions. The books are clean for this run.")
    else:
        _add_body(doc, f"{len(exceptions)} open exception(s). Showing the {len(open_rows)} most recent.")
        table = doc.add_table(rows=1 + len(open_rows), cols=5)
        table.style = "Table Grid"
        headers = ("Payment ID", "Issue", "Priority", "Amount (₹)", "Status")
        for col, header in enumerate(headers):
            _set_cell_text(table.rows[0].cells[col], header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), fill="0B3A6E", size=9)
        for index, item in enumerate(open_rows, start=1):
            fill = "FFF7F7" if index % 2 else "FFFFFF"
            _set_cell_text(table.rows[index].cells[0], item.get("payment_id"), fill=fill, size=8)
            _set_cell_text(table.rows[index].cells[1], str(item.get("mismatch_type") or "—").replace("_", " "), fill=fill, size=8)
            _set_cell_text(table.rows[index].cells[2], item.get("priority") or "Low", fill=fill, size=8)
            _set_cell_text(table.rows[index].cells[3], f"{paise_to_rupees(item.get('amount')):,.2f}", fill=fill, size=8)
            _set_cell_text(table.rows[index].cells[4], item.get("workflow_status") or "Open", fill=fill, size=8)

    monthly = list(analytics.get("monthly") or [])[-12:]
    _add_heading(doc, "5. Earnings by month", 16)
    if not monthly:
        _add_body(doc, "No monthly series yet. Load and reconcile a batch.")
    else:
        table = doc.add_table(rows=1 + len(monthly), cols=4)
        table.style = "Table Grid"
        for col, header in enumerate(("Period", "Orders", "GMV (₹)", "Net (₹)")):
            _set_cell_text(table.rows[0].cells[col], header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), fill="0D4FFF", size=9)
        for index, item in enumerate(monthly, start=1):
            fill = "F8FBFF" if index % 2 else "FFFFFF"
            _set_cell_text(table.rows[index].cells[0], item.get("label"), fill=fill, size=9)
            _set_cell_text(table.rows[index].cells[1], item.get("orders"), fill=fill, size=9)
            _set_cell_text(table.rows[index].cells[2], f"{float(item.get('gross') or 0):,.2f}", fill=fill, size=9)
            _set_cell_text(table.rows[index].cells[3], f"{float(item.get('net') or 0):,.2f}", fill=fill, size=9)

    _add_heading(doc, "6. Recent audit trail", 16)
    recent = list(audit)[:20]
    if not recent:
        _add_body(doc, "No audit events stored for this session.")
    else:
        table = doc.add_table(rows=1 + len(recent), cols=4)
        table.style = "Table Grid"
        for col, header in enumerate(("Time", "Action", "Record", "Detail")):
            _set_cell_text(table.rows[0].cells[col], header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), fill="0B3A6E", size=9)
        for index, item in enumerate(recent, start=1):
            fill = "F8FBFF" if index % 2 else "FFFFFF"
            _set_cell_text(table.rows[index].cells[0], str(item.get("timestamp") or "")[:19].replace("T", " "), fill=fill, size=8)
            _set_cell_text(table.rows[index].cells[1], item.get("action_type"), fill=fill, size=8)
            _set_cell_text(table.rows[index].cells[2], item.get("record_ids"), fill=fill, size=8)
            _set_cell_text(table.rows[index].cells[3], (item.get("details") or "")[:120], fill=fill, size=8)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    run = note.add_run(
        "Disclaimer. Missing settlements are never auto-fixed. This file is a controller snapshot of the current "
        "synthetic Razorpay batch — not a filed GST return or a bank advice."
    )
    _set_run(run, size=9, color=MUTED)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _excel_fill(hex_color: str):
    from openpyxl.styles import PatternFill
    return PatternFill("solid", fgColor=hex_color)


def build_excel_report(frame, *, date_label: str, metrics=None) -> bytes:
    """Highlighted .xlsx of a day's books. Opens in Excel or Google Sheets."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, Side
    from openpyxl.utils import get_column_letter

    metrics = metrics or {}
    book = Workbook()
    summary = book.active
    summary.title = "Summary"
    header_font = Font(bold=True, color="FFFFFF", name="Calibri")
    header_fill = _excel_fill("0B3A6E")
    thin = Border(
        left=Side(style="thin", color="D0D7E2"),
        right=Side(style="thin", color="D0D7E2"),
        top=Side(style="thin", color="D0D7E2"),
        bottom=Side(style="thin", color="D0D7E2"),
    )

    summary["A1"] = "Razor-AI daily spreadsheet"
    summary["A1"].font = Font(bold=True, size=16, color="0B3A6E")
    summary["A2"] = date_label
    kpis = [
        ("Payments", int(len(frame))),
        ("Matched", int(metrics.get("matched") or 0)),
        ("Exceptions", int(metrics.get("exceptions") or 0)),
        ("GMV (INR)", paise_to_rupees(frame["amount"].sum()) if not frame.empty and "amount" in frame.columns else 0),
        ("Fees (INR)", paise_to_rupees(frame["fee"].sum()) if not frame.empty and "fee" in frame.columns else 0),
        ("GST (INR)", paise_to_rupees(frame["tax"].sum()) if not frame.empty and "tax" in frame.columns else 0),
        ("Amount at risk (INR)", paise_to_rupees(
            frame.loc[frame["reconciliation_status"] == "exception", "amount"].sum()
        ) if not frame.empty and "reconciliation_status" in frame.columns else 0),
    ]
    summary["A4"] = "Metric"
    summary["B4"] = "Value"
    for cell in (summary["A4"], summary["B4"]):
        cell.font = header_font
        cell.fill = header_fill
    for index, (label, value) in enumerate(kpis, start=5):
        summary[f"A{index}"] = label
        summary[f"B{index}"] = value
        if "at risk" in label.lower() or label == "Exceptions":
            summary[f"B{index}"].fill = _excel_fill("FEE2E2")
        elif label == "Matched":
            summary[f"B{index}"].fill = _excel_fill("DCFCE7")
    summary["A13"] = "Red = exception / amount at risk. Green = matched. Yellow = GST or fee mismatch. This file is Excel; upload it to Google Sheets if you want it there."
    summary["A13"].font = Font(italic=True, color="64748B", size=9)
    summary.column_dimensions["A"].width = 28
    summary.column_dimensions["B"].width = 22

    sheet = book.create_sheet("Payments")
    columns = [
        "payment_id", "created_at", "amount_rupees", "fee_rupees", "tax_rupees",
        "settlement_amount_rupees", "reconciliation_status", "mismatch_type", "priority", "delta_rupees",
    ]
    titles = [
        "Payment ID", "Captured", "GMV (INR)", "Fee (INR)", "GST (INR)",
        "Credited (INR)", "Match status", "Issue", "Priority", "Delta (INR)",
    ]
    for col, title in enumerate(titles, start=1):
        cell = sheet.cell(1, col, title)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin

    rows = frame.copy() if frame is not None else None
    if rows is None or rows.empty:
        sheet["A2"] = "No payments in this date range."
    else:
        for offset, (_, row) in enumerate(rows.iterrows(), start=2):
            status = str(row.get("reconciliation_status") or "")
            mismatch = str(row.get("mismatch_type") or "")
            priority = str(row.get("priority") or "")
            values = [
                str(row.get("payment_id") or ""),
                str(row.get("created_at") or ""),
                paise_to_rupees(row.get("amount")),
                paise_to_rupees(row.get("fee")),
                paise_to_rupees(row.get("tax")),
                paise_to_rupees(row.get("settlement_amount")),
                status,
                mismatch.replace("_", " ") if mismatch and mismatch != "None" else "",
                priority,
                paise_to_rupees(row.get("delta")) if row.get("delta") is not None and str(row.get("delta")) != "nan" else "",
            ]
            fill = None
            font_color = "0F172A"
            if status == "exception":
                fill = "FECACA" if priority == "Critical" else "FED7AA"
                font_color = "7F1D1D"
            elif status == "matched":
                fill = "DCFCE7"
            if mismatch in {"tax_line_mismatch", "fee_miscalculation"}:
                fill = "FEF08A"
                font_color = "713F12"
            for col, value in enumerate(values, start=1):
                cell = sheet.cell(offset, col, value)
                cell.border = thin
                cell.font = Font(color=font_color, name="Calibri", size=10)
                if fill:
                    cell.fill = _excel_fill(fill)
        for col in range(1, len(titles) + 1):
            sheet.column_dimensions[get_column_letter(col)].width = 18
        sheet.auto_filter.ref = f"A1:{get_column_letter(len(titles))}{len(rows) + 1}"
        sheet.freeze_panes = "A2"

    legend = book.create_sheet("Legend")
    legend["A1"] = "Highlight key"
    legend["A1"].font = Font(bold=True)
    legend["A3"] = "Matched"
    legend["A3"].fill = _excel_fill("DCFCE7")
    legend["A4"] = "Exception"
    legend["A4"].fill = _excel_fill("FED7AA")
    legend["A5"] = "Critical exception"
    legend["A5"].fill = _excel_fill("FECACA")
    legend["A6"] = "Fee or GST mismatch"
    legend["A6"].fill = _excel_fill("FEF08A")

    buffer = BytesIO()
    book.save(buffer)
    return buffer.getvalue()
